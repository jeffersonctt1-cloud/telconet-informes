import sys, json, base64 as b64mod, os, io
from docx import Document
from docx.shared import Cm

_TMP_DIR = os.path.dirname(os.path.abspath(__file__))

def set_cell_text(cell, text):
    for para in cell.paragraphs:
        for run in para.runs:
            run.text = ''
    if cell.paragraphs:
        cell.paragraphs[0].add_run(str(text or ''))

def add_image_to_cell(cell, img_b64):
    if not img_b64:
        return
    if ',' in img_b64:
        img_b64 = img_b64.split(',')[1]
    img_bytes = b64mod.b64decode(img_b64)
    try:
        from PIL import Image as PILImage
        pil_img = PILImage.open(io.BytesIO(img_bytes))
        pil_img.thumbnail((800, 600), PILImage.LANCZOS)
        buf = io.BytesIO()
        pil_img.save(buf, format='JPEG', quality=75)
        img_bytes = buf.getvalue()
    except:
        pass
    tmp = os.path.join(_TMP_DIR, 'g.jpg')
    with open(tmp, 'wb') as f:
        f.write(img_bytes)
    try:
        para = cell.paragraphs[0]
        run = para.add_run()
        run.add_picture(tmp, width=Cm(5), height=Cm(4.5))
    except:
        pass
    finally:
        if os.path.exists(tmp):
            os.remove(tmp)

def fill_template_gis(data_file, output_path, template_path):
    with open(data_file, 'r', encoding='utf-8') as f:
        data = json.load(f)

    doc = Document(template_path)

    t0 = doc.tables[0]
    set_cell_text(t0.rows[0].cells[1], data.get('realizado',''))

    set_cell_text(doc.tables[1].rows[0].cells[0], data.get('nombre_punto',''))
    set_cell_text(doc.tables[2].rows[0].cells[0], data.get('nombre_cliente',''))
    set_cell_text(doc.tables[3].rows[0].cells[0], data.get('direccion',''))
    set_cell_text(doc.tables[4].rows[0].cells[0], data.get('login',''))
    set_cell_text(doc.tables[5].rows[0].cells[0], data.get('contacto',''))
    set_cell_text(doc.tables[6].rows[0].cells[0], data.get('empresa',''))

    set_cell_text(doc.tables[7].rows[0].cells[0], data.get('tecnico',''))
    if len(doc.tables[7].rows[0].cells) > 1:
        set_cell_text(doc.tables[7].rows[0].cells[1], data.get('zona',''))

    if len(doc.tables[8].rows[0].cells) > 1:
        set_cell_text(doc.tables[8].rows[0].cells[1], data.get('descripcion',''))

    t9 = doc.tables[9]
    if len(t9.rows) > 1:
        row = t9.rows[1]
        vals = ['nodo_inicio','nodo_final','ruta','caja','buffer','hilos','codigo_caja']
        for i, key in enumerate(vals):
            if i < len(row.cells):
                set_cell_text(row.cells[i], data.get(key,''))

    fotos = data.get('fotos', [])
    all_tables = list(doc.tables)

    fotos_con_img = [f for f in fotos if f.get('img') and f['img'] != '']
    n_fotos = len(fotos_con_img)

    foto_slots = [
        (10, 0), (12, 0), (12, 1), (14, 0), (14, 1),
        (16, 0), (16, 1), (18, 0), (18, 1), (20, 0),
    ]

    for idx, (t_i, col) in enumerate(foto_slots):
        if t_i < len(all_tables):
            cell = all_tables[t_i].rows[0].cells[col]
            if idx < n_fotos:
                add_image_to_cell(cell, fotos_con_img[idx]['img'])

    groups = [
        (11, 12, 1), (13, 14, 3), (15, 16, 5), (17, 18, 7), (19, 20, 9),
    ]

    if n_fotos == 0 and 10 < len(all_tables):
        tbl = all_tables[10]._tbl
        tbl.getparent().remove(tbl)

    for (lbl_t, img_t, min_idx) in reversed(groups):
        if n_fotos <= min_idx:
            for t_i in [img_t, lbl_t]:
                if t_i < len(all_tables):
                    tbl = all_tables[t_i]._tbl
                    parent = tbl.getparent()
                    if parent is not None:
                        parent.remove(tbl)

    if n_fotos > 10:
        from docx.oxml import OxmlElement as _OE
        from docx.oxml.ns import qn as _qn
        remaining = fotos_con_img[10:]
        i = 0
        while i < len(remaining):
            tr = _OE('w:tr')
            for col in range(2):
                tc = _OE('w:tc')
                tcp = _OE('w:tcPr')
                tcw = _OE('w:tcW'); tcw.set(_qn('w:w'),'4627'); tcw.set(_qn('w:type'),'dxa')
                tcp.append(tcw); tc.append(tcp)
                p = _OE('w:p'); tc.append(p)
                tr.append(tc)
            last_photo_tbl = None
            for t in doc.tables:
                last_photo_tbl = t
            if last_photo_tbl:
                last_photo_tbl._tbl.append(tr)
                tcs = tr.findall(_qn('w:tc'))
                for col_idx in range(min(2, len(remaining)-i)):
                    from docx.table import _Cell
                    cell = _Cell(tcs[col_idx], last_photo_tbl)
                    if remaining[i+col_idx].get('img'):
                        add_image_to_cell(cell, remaining[i+col_idx]['img'])
            i += 2

    t21 = all_tables[21]
    materiales = data.get('materiales', [])
    material_keys = ['fibra','canaleta','manguera_anillada','manguera_plastica','tubo_conduit',
                     'grapas_plasticas','grapas_metalicas','tacos','alambre','amarras',
                     'caja_bmx','caja_multimedia','pachtcord_fibra','patchcord_utp',
                     'duplex','simplex','tubos_fusion','sliders','abrazaderas','broca','minimanga']
    for i, key in enumerate(material_keys):
        row_idx = i + 1
        if row_idx < len(t21.rows):
            mat = next((m for m in materiales if m.get('key') == key), {})
            set_cell_text(t21.rows[row_idx].cells[1], mat.get('cantidad',''))
            set_cell_text(t21.rows[row_idx].cells[2], mat.get('medida',''))
            set_cell_text(t21.rows[row_idx].cells[3], mat.get('metros',''))
            set_cell_text(t21.rows[row_idx].cells[4], mat.get('total',''))

    from docx.oxml.ns import qn as _qn
    body = doc.element.body
    children = list(body)
    for idx, child in enumerate(children):
        tag = child.tag.split('}')[-1]
        if tag == 'p':
            from docx.text.paragraph import Paragraph
            p = Paragraph(child, doc)
            if 'Materiales Adicionales' in p.text:
                for j in range(idx+1, min(idx+5, len(children))):
                    if children[j].tag.split('}')[-1] == 'tbl':
                        from docx.table import Table
                        t_adicional = Table(children[j], doc)
                        if t_adicional.rows:
                            set_cell_text(t_adicional.rows[0].cells[0], data.get('mat_adicionales',''))
                        break
                break

    t22 = all_tables[22]
    if len(t22.rows) > 1:
        set_cell_text(t22.rows[1].cells[0], data.get('conclusiones',''))
        if len(t22.rows[1].cells) > 1:
            set_cell_text(t22.rows[1].cells[1], data.get('tiempo',''))

    doc.save(output_path)
    print('SAVED')

if __name__ == '__main__':
    fill_template_gis(sys.argv[1], sys.argv[2], sys.argv[3])

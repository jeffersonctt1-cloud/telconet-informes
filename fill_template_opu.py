import sys, json, base64 as b64mod, os, io
from docx import Document
from docx.shared import Cm, Pt, RGBColor
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

def set_cell_text(cell, text):
    for para in cell.paragraphs:
        for run in para.runs:
            run.text = ''
    if cell.paragraphs:
        cell.paragraphs[0].add_run(str(text or ''))

def set_si_no(para_idx, answer, doc):
    """Highlight SI or NO in yellow transparent for a given paragraph index"""
    p = doc.paragraphs[para_idx]
    for run in p.runs:
        txt = run.text.strip()
        rPr = run._r.find(qn('w:rPr'))
        if rPr is None:
            rPr = OxmlElement('w:rPr')
            run._r.insert(0, rPr)
        # Remove existing highlight
        hl = rPr.find(qn('w:highlight'))
        if hl is not None:
            rPr.remove(hl)
        # Apply yellow highlight to selected option
        if (answer == 'SI' and txt == 'SI') or (answer == 'NO' and txt == 'NO'):
            hl = OxmlElement('w:highlight')
            hl.set(qn('w:val'), 'yellow')
            rPr.append(hl)

def add_signature_to_cell(cell, sig_b64):
    """Add signature image to cell"""
    if not sig_b64 or len(sig_b64) < 100:
        return
    try:
        if ',' in sig_b64:
            sig_b64 = sig_b64.split(',')[1]
    except:
        return
    img_bytes = b64mod.b64decode(sig_b64)
    try:
        from PIL import Image as PILImage
        pil_img = PILImage.open(io.BytesIO(img_bytes))
        pil_img.thumbnail((400, 200), PILImage.LANCZOS)
        buf = io.BytesIO()
        pil_img.save(buf, format='PNG')
        img_bytes = buf.getvalue()
    except:
        pass
    tmp = 'C:\\t\\sig.png'
    os.makedirs('C:\\t', exist_ok=True)
    with open(tmp, 'wb') as f:
        f.write(img_bytes)
    try:
        para = cell.paragraphs[0]
        run = para.add_run()
        run.add_picture(tmp, width=Cm(5), height=Cm(2))
    except:
        pass
    finally:
        if os.path.exists(tmp):
            os.remove(tmp)

def fill_template_opu(data_file, output_path, template_path):
    with open(data_file, 'r', encoding='utf-8') as f:
        data = json.load(f)

    doc = Document(template_path)

    # TABLE 0: Datos generales
    t0 = doc.tables[0]
    set_cell_text(t0.rows[0].cells[1], data.get('fecha', ''))       # FECHA auto
    set_cell_text(t0.rows[1].cells[1], data.get('tecnico', ''))      # NOMBRE TÉCNICO
    set_cell_text(t0.rows[2].cells[1], data.get('instalacion', ''))  # INSTALACIÓN
    set_cell_text(t0.rows[3].cells[1], data.get('cliente', ''))      # CLIENTE
    set_cell_text(t0.rows[4].cells[1], data.get('login', ''))        # LOGIN
    set_cell_text(t0.rows[5].cells[1], data.get('fibra', ''))        # FIBRA
    set_cell_text(t0.rows[6].cells[1], data.get('cable_acometida',''))# CABLE ACOMETIDA
    set_cell_text(t0.rows[7].cells[1], data.get('codigo_caja', ''))  # CÓDIGO CAJA
    set_cell_text(t0.rows[8].cells[1], data.get('hilo', ''))         # HILO
    set_cell_text(t0.rows[9].cells[1], data.get('sw', ''))           # SW
    set_cell_text(t0.rows[10].cells[1], data.get('puerto', ''))      # PUERTO
    set_cell_text(t0.rows[11].cells[1], data.get('responsable', '')) # RESPONSABLE IAC

    # TABLE 1: Materiales (col 1 = cantidad1, col 2 = cantidad2)
    t1 = doc.tables[1]
    mat_map = [
        ('fibra_m', 0), ('cable_utp', 1), ('conector_rj45', 2),
        ('pathcord_fo', 3), ('simplex', 4), ('botitas', 5),
        ('tubillos', 6), ('duplex', 7), ('caja_multimedia', 8),
        ('serie_tx_nodo', 9), ('serie_tx_cliente', 10),
        ('router_modelo', 11), ('router_serie', 12),
    ]
    for key, row_idx in mat_map:
        if row_idx < len(t1.rows):
            val = data.get(key, '')
            if val:
                set_cell_text(t1.rows[row_idx].cells[1], str(val))

    # Materiales extras (rows 14,15,16 = ganchos, amarras, pinzas)
    extras = [('ganchos', 14), ('amarras', 15), ('pinzas', 16)]
    for key, row_idx in extras:
        if row_idx < len(t1.rows):
            val = data.get(key, '')
            if val:
                set_cell_text(t1.rows[row_idx].cells[2], str(val))

    # SI/NO questions — paragraph indices: 25,27,29,31,33,35
    sinos = data.get('sinos', [])
    si_no_paras = [25, 27, 29, 31, 33, 35]
    for i, para_idx in enumerate(si_no_paras):
        answer = sinos[i] if i < len(sinos) else ''
        if answer:
            set_si_no(para_idx, answer, doc)

    # TABLE 2: Firmas (2 cells: técnico, cliente)
    t2 = doc.tables[2]
    sig_tecnico = data.get('firma_tecnico', '')
    sig_cliente = data.get('firma_cliente', '')
    if sig_tecnico:
        add_signature_to_cell(t2.rows[0].cells[0], sig_tecnico)
    if sig_cliente:
        add_signature_to_cell(t2.rows[0].cells[1], sig_cliente)

    doc.save(output_path)
    print('SAVED')

if __name__ == '__main__':
    fill_template_opu(sys.argv[1], sys.argv[2], sys.argv[3])

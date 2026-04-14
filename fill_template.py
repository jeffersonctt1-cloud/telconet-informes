import sys, json, base64 as b64mod, os, io
from docx import Document
from docx.shared import Cm

_TMP_DIR = os.path.dirname(os.path.abspath(__file__))

def set_cell_text(cell, text):
    for para in cell.paragraphs:
        for run in para.runs:
            run.text = ''
    if cell.paragraphs:
        cell.paragraphs[0].add_run(str(text or ''))

def set_label_value(cell, value):
    for para in cell.paragraphs:
        if para.text.strip():
            for run in para.runs[1:]:
                run.text = ''
            if para.runs:
                para.runs[0].text = para.runs[0].text.rstrip() + ' '
            para.add_run(str(value or ''))
            return

def set_photo_label_cell(cell, fig_num, coords):
    paras = cell.paragraphs
    if len(paras) > 0:
        for run in paras[0].runs:
            run.text = ''
        if paras[0].runs:
            paras[0].runs[0].text = f'Fig {fig_num}.  '
        else:
            paras[0].add_run(f'Fig {fig_num}.  ')
    if len(paras) > 1:
        for run in paras[1].runs:
            run.text = ''
        if paras[1].runs:
            paras[1].runs[0].text = f'Coordenadas: {coords}'
        else:
            paras[1].add_run(f'Coordenadas: {coords}')

def clear_photo_label_cell(cell):
    for para in cell.paragraphs:
        for run in para.runs:
            run.text = ''

def add_image_to_cell(cell, img_b64):
    if not img_b64:
        return
    if ',' in img_b64:
        img_b64 = img_b64.split(',')[1]
    img_bytes = b64mod.b64decode(img_b64)
    try:
        from PIL import Image as PILImage
        pil_img = PILImage.open(io.BytesIO(img_bytes))
        pil_img.thumbnail((800, 600), PILImage.LANCZOS)
        buf = io.BytesIO()
        pil_img.save(buf, format='JPEG', quality=75)
        img_bytes = buf.getvalue()
    except:
        pass
    tmp = os.path.join(_TMP_DIR, 'p.jpg')
    with open(tmp, 'wb') as f:
        f.write(img_bytes)
    try:
        para = cell.paragraphs[0]
        run = para.add_run()
        run.add_picture(tmp, width=Cm(5), height=Cm(4.5))
    except:
        pass
    finally:
        if os.path.exists(tmp):
            os.remove(tmp)

def remove_row(table, row):
    table._tbl.remove(row._tr)

def fill_template(data_file, output_path, template_path):
    with open(data_file, 'r', encoding='utf-8') as f:
        data = json.load(f)

    doc = Document(template_path)

    # TABLE 0: Encabezado
    t0 = doc.tables[0]
    set_label_value(t0.rows[1].cells[0], data.get('realizado',''))
    set_label_value(t0.rows[1].cells[1], data.get('cargo',''))
    set_label_value(t0.rows[2].cells[0], data.get('fecha',''))
    set_label_value(t0.rows[2].cells[1], data.get('dpto',''))

    # TABLE 1: Tarea / Entidad
    t1 = doc.tables[1]
    set_cell_text(t1.rows[0].cells[1], data.get('tarea',''))
    set_cell_text(t1.rows[1].cells[1], data.get('entidad','CNT'))

    # TABLE 2: Datos generales
    t2 = doc.tables[2]
    set_cell_text(t2.rows[1].cells[1], data.get('proyecto',''))
    set_cell_text(t2.rows[2].cells[1], data.get('parroquia',''))
    set_cell_text(t2.rows[2].cells[3], data.get('ciudad',''))
    set_cell_text(t2.rows[3].cells[1], data.get('dir',''))

    # TABLE 3: Pozos
    t3 = doc.tables[3]
    max_rows = len(t3.rows) - 1
    wells = [w for w in data.get('wells',[]) if w.get('codigo') or w.get('dir') or w.get('coords')]
    rows_to_remove = []
    for i in range(max_rows):
        row = t3.rows[i + 1]
        if i < len(wells):
            w = wells[i]
            set_cell_text(row.cells[0], str(i + 1))
            set_cell_text(row.cells[1], w.get('codigo',''))
            set_cell_text(row.cells[2], w.get('dir',''))
            coords_raw = w.get('coords','')
            lat, lon = '', ''
            if coords_raw and ',' in coords_raw:
                parts = coords_raw.split(',', 1)
                lat = parts[0].strip()
                lon = parts[1].strip()
            else:
                lat = coords_raw
            set_cell_text(row.cells[3], lat)
            set_cell_text(row.cells[4], lon)
            set_cell_text(row.cells[5], w.get('estado',''))
            set_cell_text(row.cells[6], w.get('autor',''))
        else:
            rows_to_remove.append(row)
    for row in rows_to_remove:
        remove_row(t3, row)

    # TABLE 4: Fotos
    t4 = doc.tables[4]
    fotos = data.get('fotos', [])
    n = len(fotos)
    all_rows = list(t4.rows)

    label_slots = [(1,0,1),(1,1,2),(3,0,3),(3,1,4),(5,0,5)]
    for (ri, col, fig_n) in label_slots:
        if ri < len(all_rows):
            cell = all_rows[ri].cells[col]
            idx = fig_n - 1
            if idx < n:
                set_photo_label_cell(cell, fig_n, fotos[idx].get('coords',''))
            else:
                clear_photo_label_cell(cell)

    img_slots = [(0,0,0),(0,1,1),(2,0,2),(2,1,3),(4,0,4)]
    for (ri, col, idx) in img_slots:
        if ri < len(all_rows):
            cell = all_rows[ri].cells[col]
            if idx < n and fotos[idx].get('img'):
                add_image_to_cell(cell, fotos[idx]['img'])

    rows_needed = set()
    if n >= 1: rows_needed.update([0, 1])
    if n >= 3: rows_needed.update([2, 3])
    if n >= 5: rows_needed.update([4, 5])

    for r_idx in range(len(all_rows)-1, -1, -1):
        if r_idx not in rows_needed:
            remove_row(t4, all_rows[r_idx])

    # Calcular cuántas filas se eliminaron de t4
    # Con 5 fotos: 6 filas → 0 eliminadas → 13 párrafos vacíos OK
    # Con 3-4 fotos: 4 filas → 2 eliminadas → eliminar ~4 párrafos vacíos
    # Con 1-2 fotos: 2 filas → 4 eliminadas → eliminar ~8 párrafos vacíos
    # Con 0 fotos: 0 filas → 6 eliminadas → eliminar todos los párrafos vacíos
    filas_eliminadas = len(all_rows) - len(rows_needed)

    # Eliminar párrafos vacíos sobrantes después de t4
    # Cada par de filas eliminadas equivale a ~4 párrafos vacíos sobrantes
    # Dejamos siempre 1 párrafo vacío de separación
    from docx.text.paragraph import Paragraph as _Para
    from docx.oxml.ns import qn as _qn

    t4_elem = doc.tables[4]._tbl
    body = doc.element.body
    body_children = list(body)

    # Encontrar posición de t4 en el body
    t4_pos = None
    for idx, child in enumerate(body_children):
        if child is t4_elem:
            t4_pos = idx
            break

    if t4_pos is not None and filas_eliminadas > 0:
        # Contar párrafos vacíos consecutivos después de t4
        paras_vacios = []
        for child in body_children[t4_pos + 1:]:
            tag = child.tag.split('}')[-1]
            if tag == 'p':
                p = _Para(child, doc)
                if p.text.strip() == '' and len(p.runs) == 0:
                    paras_vacios.append(child)
                else:
                    break  # encontramos contenido real, parar
            else:
                break  # encontramos tabla u otro elemento, parar

        # Calcular cuántos eliminar: cada 2 filas eliminadas = ~4 párrafos sobrantes
        # Máximo: dejar siempre 1 párrafo de separación
        n_eliminar = min(filas_eliminadas * 2, max(0, len(paras_vacios) - 1))

        for p_elem in paras_vacios[:n_eliminar]:
            body.remove(p_elem)

    # TABLE 5: Observaciones
    t5 = doc.tables[5]
    set_cell_text(t5.rows[0].cells[0], data.get('obs',''))

    # TABLE 6: Recomendaciones
    t6 = doc.tables[6]
    set_cell_text(t6.rows[0].cells[0], data.get('rec',''))

    doc.save(output_path)
    print('SAVED')

if __name__ == '__main__':
    fill_template(sys.argv[1], sys.argv[2], sys.argv[3])
